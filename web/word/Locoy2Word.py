import os
import re
from functools import reduce

import cell as cell
import requests
import xlwt
from bs4 import BeautifulSoup as bs
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches
from requests.cookies import RequestsCookieJar

from web.DownloadImage import download_img, get_img_url

user_agent = 'Mozilla/5.0 (Windows NT 6.1; Win64; x64; Trident/7.0; rv:11.0) like Gecko'

header = {
    "User-Agent": user_agent
}

cookies = '1pqA_2132_saltkey=efXchfXC; 1pqA_2132_lastvisit=1561599634;' \
          ' 1pqA_2132_lastact=1561603263%09plugin.php%09; 1pqA_2132_sendmail=1;' \
          ' 1pqA_2132_ulastactivity=1561603262%7C0; 1pqA_2132_auth=dd38ojIxbrKd5q0' \
          'y9b%2FO4a0HkIHYp8Otonhtm5wB9znZXbM4QgPyoS6osZPiY2TTa9hLI%2BbVMWCbAF8rtXFdp' \
          '9nQsbeK; 1pqA_2132_lastcheckfeed=1290934%7C1561603262; 1pqA_2132_check' \
          'follow=1; 1pqA_2132_lip=110.178.211.35%2C1561603262'
d = dict(map(lambda x: x.split('='), cookies.split(';')))
# 添加登录需要的cookie
cookie_jar = RequestsCookieJar()
for k, v in d.items():
    cookie_jar.set(k, v)

base_url = 'https://bbs.shiyebian.org'

# 获取题的解析
question_result = '/plugin.php?id=kaoshi:showbox&from=exam&mod=showcontent' \
                  '&infloat=yes&handlekey=showbox&inajax=1&ajaxtarget=fwin_content_showbox&cid='
# 历年真题
true_exam = "/plugin.php?id=kaoshi:list&tid=92&sid=79&mode=2&page="

# 真题
exam = "/plugin.php?id=kaoshi:paper&sid=79&pid="

# 模块特训
model_chapter = '/plugin.php?id=kaoshi:chapter&sid=79'

# 模块下的分类
model_list = '/plugin.php?id=kaoshi:list&sid=79&cid=62'

# 模块下的分类下的题
model_paper = '/plugin.php?id=kaoshi:paper&sid=79&pid=11249&cid=62'


# 启动爬虫
def start_spider(list_url):
    # for i in range(1, 18):
    #     if i == 2:
    #         return
    #     req = requests.get(base_url + true_exam + str(i), headers=header, cookies=cookie_jar)
    #     soup = bs(str(req.content, 'GBK'), 'html.parser')
    #     for data in re.findall(r'pid=([0-9]*)', str(soup.find('ul', class_='ul-li'))):
    #         get_bank_list(list_url + str(data))
    get_bank_list(list_url + str(23885))


def get_bank_list(list_url):
    req = requests.get(list_url, headers=header, cookies=cookie_jar)
    if req.status_code == 200:
        soup = bs(str(req.content, 'GBK'), 'html.parser')
        # <div class="paper-bt">2019年浙江省公务员录用考试《行测》真题及解析（B类）</div>
        exam_name = soup.find_all('div', class_='paper-bt')[0].get_text()
        create_excel(exam_name, soup)


# 写入excel
def create_excel(exam_name, soup):
    os.makedirs('./data/', exist_ok=True)
    document = Document()
    document.styles['Normal'].font.size = Pt(16)
    document.styles['Normal'].font.name = u'宋体'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    # document.add_picture('200.PNG')
    # document.add_paragraph('my paragraphfdjakhfkjdhgf翻江倒海房间爱货到付款交换空间安顿好副科级哈的控件和父控件哈卡交话费看就好')

    write_data(exam_name, soup, document)
    document.save(exam_name + '.docx')


# 写入数据
def write_data(exam_name, soup, document):
    for i, value in enumerate(soup.find_all('div', class_='exam-box')):
        print('正在写入 eid = %s 的问题' % value.map['eid'])
        print(str(i))
        # 题干(有图)
        content_run = document.add_paragraph().add_run()
        for tag in value.find('div', class_='exam-subject').contents:
            if str(tag).startswith('<img'):
                content_run.add_picture(download_img(get_img_url(tag)))
            else:
                content_run.add_text(str(tag))

        # 选项(可能有图)
        option_labels = value.find('ul', class_='exam-options').find_all('label')
        for label in option_labels:
            option_run = document.add_paragraph().add_run()
            option_run.add_text(label.get_text())
            for img_tag in label.find_all('img'):
                option_run.add_picture(download_img(get_img_url(img_tag)))

        # 答案(纯文本)
        answer = value.find('span', class_='result').get_text()
        document.add_paragraph("【答案】：")
        document.add_paragraph(answer)

        # 解析(有图)
        # 题id
        document.add_paragraph("【解析】：")
        eid = value.map['eid']
        # 单独创建请求获取解析
        req = requests.get(base_url + question_result + eid, headers=header, cookies=cookie_jar)
        if req.status_code == 200:
            soup = bs(replace_cdata(str(req.content, 'GBK')), 'html.parser')
            for data in soup.find_all('p'):
                if str(data).find('【解析】') > 0:
                    content_run = document.add_paragraph().add_run()
                    for tag in data.contents:
                        if str(tag).startswith('<img'):
                            content_run.add_picture(download_img(get_img_url(tag)))
                        elif str(tag).startswith('<br/'):
                            document.add_paragraph(" ")
                        elif str(tag).startswith('<b>'):
                            pass
                        else:
                            content_run.add_text(str(tag))

        # 跟下一题隔开
        document.add_page_break()
        print('写入完成 eid = %s 的问题' % value.map['eid'])


# 替换html标签
def replace_html_tag(s):
    result = s.replace('<p>', '') \
        .replace('</p>', '') \
        .replace('<br/>', '\r\n')
    # <u>替换下划线</u>
    return re.sub(re.compile(r"<u.*?</u>", re.S), "_______", result)


def replace_cdata(s):
    return s.replace("<![CDATA[", "").replace(']]>', '')


if __name__ == '__main__':
    start_spider(base_url + exam)
